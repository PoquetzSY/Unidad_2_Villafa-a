from docx import Document

def word_f(texto):
    doc = Document()
    doc.add_paragraph(texto)
    doc.save("Prueba.docx")

    doc = Document("Prueba.docx")
    for paragraph in doc.paragraphs:
        print(paragraph.text)